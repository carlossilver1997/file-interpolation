import os
import fpdf
from utils.util import calculateAge
from docx import Document
from docx.text.paragraph import Paragraph
from decouple import config
import smtplib
from os.path import basename
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

def read_and_interpolate_file(route, person, ext):
    try:
        with open(route) as f:
            content = f.read()
            content = content.replace("[NAME]", person["givenNames"].title()).replace("[LAST_NAME_1]", person["lastName1"].title()).replace("[LAST_NAME_2]", person["lastName2"].title()).replace("[COMPANY]", person["company"].title()).replace("[POSITION]", person["position"].upper()).replace("[BIRTH_DATE]", person["birthDate"]).replace("[PHONE]", person["phoneNumber"]).replace("[EMAIL]", person["email"]).replace("[ADDRESS]", person["address"].title()).replace("[AGE]", str(calculateAge(person["birthDate"])))            
        cwd = os.getcwd()
        lines = content.split('\n')
        if ext == 'pdf':
            pdf = fpdf.FPDF(format='letter') #pdf format
            pdf.add_page() #create new page
            pdf.set_font("Arial", size=12) # font and textsize
            for index, line in enumerate(lines):
                pdf.cell(200, 10, txt=line, ln=(index+1), align='L')
            pdf.output("{}/static/{}.{}".format(cwd,person["id"],ext))
            return
        if ext == 'docx':
            document = Document()
            for index, line in enumerate(lines):
                if index == 0:
                    document.add_heading(line, 0)
                else:
                    document.add_paragraph(line)
            document.save("{}/static/{}.{}".format(cwd,person["id"],ext))
            return
        with open('{}/static/{}.txt'.format(cwd,person["id"]), 'w') as file:
            file.write(content)

    except FileNotFoundError as fne:
        raise fne
    except FileExistsError as fee:
        raise fne

def sendEmail(background_tasks, route, person):
    try:
        with open(route) as f:
            content = f.read()
            content = content.replace("[NAME]", person["givenNames"].title()).replace("[LAST_NAME_1]", person["lastName1"].title()).replace("[LAST_NAME_2]", person["lastName2"].title()).replace("[COMPANY]", person["company"].title()).replace("[POSITION]", person["position"].upper()).replace("[BIRTH_DATE]", person["birthDate"]).replace("[PHONE]", person["phoneNumber"]).replace("[EMAIL]", person["email"]).replace("[ADDRESS]", person["address"].title()).replace("[AGE]", str(calculateAge(person["birthDate"])))            
        cwd = os.getcwd()
        lines = content.split('\n')
        pdf = fpdf.FPDF(format='letter') #pdf format
        pdf.add_page() #create new page
        pdf.set_font("Arial", size=12) # font and textsize
        for index, line in enumerate(lines):
            pdf.cell(200, 10, txt=line, ln=(index+1), align='L')
        pdf.output("{}/static/{}.pdf".format(cwd,person["id"]))
        document = Document()
        for index, line in enumerate(lines):
            if index == 0:
                document.add_heading(line, 0)
            else:
                document.add_paragraph(line)
        document.save("{}/static/{}.docx".format(cwd,person["id"]))
        with open('{}/static/{}.txt'.format(cwd,person["id"]), 'w') as file:
            file.write(content)
        # SEND FILES WITH EXTENSION: PDF, DOCX, TXT

        
        docs = [
            '{}/static/{}.pdf'.format(cwd, person["id"]),
            '{}/static/{}.txt'.format(cwd, person["id"]),
            '{}/static/{}.docx'.format(cwd, person["id"])
        ]
        mail_content = '''Hello,
        This is a test mail.
        In this mail we are sending some attachments.
        The mail is sent using Python SMTP library.
        Thank You
        '''
        #The mail addresses and password
        sender_address = config("EMAIL")
        sender_pass = config("PASSWORD")
        receiver_address = person["email"]
        print(receiver_address)
        #Setup the MIME
        message = MIMEMultipart()
        message['From'] = sender_address
        message['To'] = receiver_address
        message['Subject'] = 'Carta agradecimiento en pdf, docx, txt'
        #The subject line
        #The body and the attachments for the mail
        for f in docs:
            with open(f, "rb") as fil:
                part = MIMEApplication(
                    fil.read(),
                    Name=basename(f)
                )
                part['Content-Disposition'] = 'attachment; filename="%s"' % basename(f)
                message.attach(part)
                #Create SMTP session for sending the mail
        session = smtplib.SMTP('smtp.gmail.com', 587) #use gmail with port
        session.starttls() #enable security
        session.login(sender_address, sender_pass) #login with mail_id and password
        text = message.as_string()
        session.sendmail(sender_address, receiver_address, text)
        session.quit()
        print('Mail Sent')

    except FileNotFoundError as fne:
        raise fne
    except FileExistsError as fee:
        raise fne
